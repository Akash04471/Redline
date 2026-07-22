import docx

def create_sample_contract():
    doc = docx.Document()
    doc.add_heading('Mutual Non-Disclosure Agreement', 0)

    doc.add_paragraph('This Mutual Non-Disclosure Agreement is entered into by the parties...')

    doc.add_heading('1. Confidentiality', level=1)
    doc.add_paragraph('Each party shall keep confidential all information disclosed by the other party and shall not disclose such information to any third party without prior written consent. This obligation shall survive for a period of 5 years.')

    doc.add_heading('2. Indemnification', level=1)
    doc.add_paragraph('Provider shall indemnify, defend, and hold harmless the Client from any claims, damages, liabilities, and expenses arising out of or related to this agreement, provided that liability is capped at $10,000.')

    doc.add_heading('3. Data Processing', level=1)
    doc.add_paragraph('The Provider agrees to process personal data strictly in accordance with GDPR. Upon request, the Provider must delete personal data within 30 days of notice.')

    doc.add_heading('4. Termination', level=1)
    doc.add_paragraph('Either party may terminate this agreement at any time by providing 10 days written notice. In the event of termination, all confidential information must be destroyed.')

    doc.add_heading('5. Governing Law', level=1)
    doc.add_paragraph('This Agreement shall be governed by and construed in accordance with the laws of the State of Delaware, without regard to its conflict of law provisions.')

    doc.save('sample_contract.docx')
    print("Saved sample_contract.docx")

if __name__ == "__main__":
    create_sample_contract()

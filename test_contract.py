import docx

doc = docx.Document()
doc.add_heading('MUTUAL NON-DISCLOSURE AGREEMENT', 0)

doc.add_paragraph('This Mutual Non-Disclosure Agreement (this "Agreement") is entered into as of July 22, 2026.')

doc.add_paragraph('1. Definition of Confidential Information. "Confidential Information" means any information disclosed by either party to the other party, either directly or indirectly, in writing, orally or by inspection of tangible objects (including without limitation documents, prototypes, samples, plant and equipment), which is designated as "Confidential," "Proprietary" or some similar designation. Information communicated orally shall be considered Confidential Information if such information is confirmed in writing as being Confidential Information within a reasonable time after the initial disclosure. Confidential Information may also include information disclosed to a disclosing party by third parties.')

doc.add_paragraph('2. Exceptions. Confidential Information shall not, however, include any information which (i) was publicly known and made generally available in the public domain prior to the time of disclosure by the disclosing party; (ii) becomes publicly known and made generally available after disclosure by the disclosing party to the receiving party through no action or inaction of the receiving party; (iii) is already in the possession of the receiving party at the time of disclosure by the disclosing party as shown by the receiving party\'s files and records immediately prior to the time of disclosure; (iv) is obtained by the receiving party from a third party without a breach of such third party\'s obligations of confidentiality; (v) is independently developed by the receiving party without use of or reference to the disclosing party\'s Confidential Information, as shown by documents and other competent evidence in the receiving party\'s possession; or (vi) is required by law to be disclosed by the receiving party, provided that the receiving party gives the disclosing party prompt written notice of such requirement prior to such disclosure and assistance in obtaining an order protecting the information from public disclosure.')

doc.add_paragraph('3. Non-Use and Non-Disclosure. Each party agrees not to use any Confidential Information of the other party for any purpose except to evaluate and engage in discussions concerning a potential business relationship between the parties. Each party agrees not to disclose any Confidential Information of the other party to third parties or to such party\'s employees, except to those employees of the receiving party who are required to have the information in order to evaluate or engage in discussions concerning the contemplated business relationship.')

doc.add_paragraph('4. Maintenance of Confidentiality. Each party agrees that it shall take reasonable measures to protect the secrecy of and avoid disclosure and unauthorized use of the Confidential Information of the other party. Without limiting the foregoing, each party shall take at least those measures that it takes to protect its own most highly confidential information and shall ensure that its employees who have access to Confidential Information of the other party have signed a non-use and non-disclosure agreement in content similar to the provisions hereof, prior to any disclosure of Confidential Information to such employees.')

doc.add_paragraph('5. No Obligation. Nothing herein shall obligate either party to proceed with any transaction between them, and each party reserves the right, in its sole discretion, to terminate the discussions contemplated by this Agreement concerning the business opportunity.')

doc.add_paragraph('6. No Warranty. ALL CONFIDENTIAL INFORMATION IS PROVIDED "AS IS." EACH PARTY MAKES NO WARRANTIES, EXPRESS, IMPLIED OR OTHERWISE, REGARDING ITS ACCURACY, COMPLETENESS OR PERFORMANCE.')

doc.add_paragraph('7. Return of Materials. All documents and other tangible objects containing or representing Confidential Information which have been disclosed by either party to the other party, and all copies thereof which are in the possession of the other party, shall be and remain the property of the disclosing party and shall be promptly returned to the disclosing party upon the disclosing party\'s written request.')

doc.add_paragraph('8. No License. Nothing in this Agreement is intended to grant any rights to either party under any patent, mask work right or copyright of the other party, nor shall this Agreement grant any party any rights in or to the Confidential Information of the other party except as expressly set forth herein.')

doc.add_paragraph('9. Term. The obligations of each receiving party hereunder shall survive until such time as all Confidential Information of the other party disclosed hereunder becomes publicly known and made generally available through no action or inaction of the receiving party.')

doc.add_paragraph('10. Remedies. Each party agrees that any violation or threatened violation of this Agreement may cause irreparable injury to the other party, entitling the other party to seek injunctive relief in addition to all legal remedies.')

doc.add_paragraph('11. Miscellaneous. This Agreement shall bind and inure to the benefit of the parties hereto and their successors and assigns. This Agreement shall be governed by the laws of the State of California, without reference to conflict of laws principles. This document contains the entire agreement between the parties with respect to the subject matter hereof, and neither party shall have any obligation, express or implied by law, with respect to trade secret or proprietary information of the other party except as set forth herein.')

doc.save('sample_contract.docx')

from src.parsing.extractor import extract_text, split_into_clauses

with open('sample_contract.docx', 'rb') as f:
    text = extract_text(f.read(), 'sample_contract.docx')

clauses = split_into_clauses(text)
for c in clauses:
    print(f'Clause {c.clause_index} [{c.clause_category}]: {c.raw_text[:70]}...')
